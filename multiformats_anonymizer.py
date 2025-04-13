import spacy
from docx import Document
import csv
import os
import pdfplumber
import pandas as pd
import fitz  # PyMuPDF

from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

# This script anonymizes sensitive information in DOCX, XLSX, and PDF files.
# It uses spaCy for entity recognition and redacts sensitive entities in PDFs.
# It also logs the changes made to the documents.

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# Define custom placeholders for each entity type
CUSTOM_LABELS = {
    "ORG": "[CLIENT_NAME]",
    "PERSON": "[PERSON]",
    "GPE": "[LOCATION]",
    "PRODUCT": "[PRODUCT_NAME]"
}

# Keep a log of replacements/redactions
replacement_log = []

def anonymize_text(text):
    """Replace named entities with custom placeholder tokens."""
    doc = nlp(text)
    new_text = text
    for ent in reversed(doc.ents):
        if ent.label_ in CUSTOM_LABELS:
            placeholder = CUSTOM_LABELS[ent.label_]
            replacement_log.append((ent.text, placeholder))
            new_text = new_text[:ent.start_char] + placeholder + new_text[ent.end_char:]
    return new_text

def save_replacement_log(log_path):
    """Write a CSV log of all anonymized or redacted terms."""
    with open(log_path, mode='w', newline='', encoding='utf-8') as log_file:
        writer = csv.writer(log_file)
        writer.writerow(["Original", "Replacement"])
        for original, replacement in replacement_log:
            writer.writerow([original, replacement])

def anonymize_docx_full(input_path, output_path, log_path):
    """Anonymize DOCX: paragraphs, tables, headers, and footers."""
    doc = Document(input_path)

    # Anonymize normal text
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                run.text = anonymize_text(run.text)

    # Anonymize tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            run.text = anonymize_text(run.text)

    # Anonymize headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    run.text = anonymize_text(run.text)
        for para in section.footer.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    run.text = anonymize_text(run.text)

    doc.save(output_path)
    save_replacement_log(log_path)
    print(f"✅ DOCX saved: {output_path}\n📝 Log saved: {log_path}")

def anonymize_excel(input_path, output_path, log_path):
    """
    Anonymize Excel cells using openpyxl to preserve formatting.
    Only string-type cells are scanned and anonymized.
    """
    # Load the Excel workbook and preserve styles
    wb = load_workbook(input_path)
    replacement_log.clear()  # Reset log before processing

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    original_value = cell.value
                    new_value = anonymize_text(original_value)
                    if new_value != original_value:
                        replacement_log.append((original_value, new_value))
                        cell.value = new_value  # Update cell with redacted value

    # Save the new Excel file
    wb.save(output_path)

    # Save redaction log
    with open(log_path, mode='w', newline='', encoding='utf-8') as log_file:
        writer = csv.writer(log_file)
        writer.writerow(["Original", "Replacement"])
        for original, replacement in replacement_log:
            writer.writerow([original, replacement])

    print(f"✅ Excel anonymized and saved to: {output_path}")
    print(f"📝 Redaction log saved to: {log_path}")

def redact_pdf(input_path, output_path, log_path):
    """
    Redact sensitive entities in a PDF using spaCy and PyMuPDF (fitz).
    This version properly applies redactions to remove underlying text.
    """

    # Open the PDF using PyMuPDF
    doc = fitz.open(input_path)

    # List to store all redacted entities for logging
    all_entities = []

    # Loop through each page in the PDF
    for page_num in range(len(doc)):
        page = doc[page_num]

        # Extract text content from the page
        text = page.get_text()

        # Use spaCy to identify sensitive entities on this page
        entities = [(ent.text, ent.label_) for ent in nlp(text).ents if ent.label_ in CUSTOM_LABELS]
        all_entities.extend(entities)

        # For each sensitive word/entity, find its location and mark for redaction
        for word, label in entities:
            areas = page.search_for(word)
            for area in areas:
                # Add a redaction box with solid black fill
                page.add_redact_annot(area, fill=(0, 0, 0))

        # 🛡️ APPLY the redactions (this step removes the text underneath!)
        page.apply_redactions()

    # Save the redacted PDF
    doc.save(output_path, garbage=4, deflate=True, clean=True)

    # Save redacted entity log
    with open(log_path, mode='w', newline='', encoding='utf-8') as log_file:
        writer = csv.writer(log_file)
        writer.writerow(["Redacted Text", "Entity Type"])
        for word, label in all_entities:
            writer.writerow([word, label])

    # Done!
    print(f"✅ PDF redacted and saved to: {output_path}")
    print(f"📝 Redaction log saved to: {log_path}")

if __name__ == "__main__":
    print("\n📂 Supported File Types: DOCX | XLSX | PDF")
    input_path = input("🔍 Enter the full path of the file to anonymize or redact: ")

    if not os.path.exists(input_path):
        print("❌ Error: File not found.")
        exit()

    filename, ext = os.path.splitext(input_path)
    ext = ext.lower()

    output_path = filename + "_anonymized" + ext
    log_path = filename + "_log.csv"

    if ext == ".docx":
        anonymize_docx_full(input_path, output_path, log_path)

    elif ext == ".xlsx":
        anonymize_excel(input_path, output_path, log_path)

    elif ext == ".pdf":
        redact_pdf(input_path, output_path, log_path)

    else:
        print("⚠️ Unsupported file type. Please use .docx, .xlsx, or .pdf only.")
