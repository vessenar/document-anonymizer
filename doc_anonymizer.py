import spacy
from docx import Document
import csv
import os
import fitz  # PyMuPDF
from datetime import datetime
from openpyxl import load_workbook

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# Define custom placeholders for each entity type
CUSTOM_LABELS = {
    "ORG": "[ORG_NAME]",
    "PERSON": "[PERSON]",
    "GPE": "[LOCATION]",
    "PRODUCT": "[PRODUCT_NAME]"
}

# Global set for unique redactions
unique_redactions = set()

# Global dictionary to store entity-to-placeholder mappings
entity_mapping = {}

# Global whitelist for keywords that should not be redacted
WHITELIST = []

def get_whitelist_from_user():
    """Prompt the user to enter a comma-separated list of keywords to whitelist."""
    custom_whitelist = input("Enter comma-separated keywords to whitelist (or leave blank): ").strip().strip("'\"")
    if custom_whitelist:
        return [term.strip() for term in custom_whitelist.split(",")]
    return []

def anonymize_text(text, filename):
    """Replace named entities with custom tokens and ensure consistent redaction."""
    doc = nlp(text)
    new_text = text
    for ent in reversed(doc.ents):
        if ent.label_ in CUSTOM_LABELS:
            # Skip redaction if the entity is in the whitelist
            if ent.text in WHITELIST:
                continue

            # Check if the entity is already in the mapping
            if ent.text not in entity_mapping:
                placeholder = CUSTOM_LABELS[ent.label_]
                entity_mapping[ent.text] = placeholder
                unique_redactions.add((filename, ent.text, placeholder))
            else:
                placeholder = entity_mapping[ent.text]

            # Replace the entity with its placeholder
            new_text = new_text[:ent.start_char] + placeholder + new_text[ent.end_char:]
    return new_text

def save_replacement_log(log_path):
    """Write a CSV log of all unique redactions with filename and timestamp."""
    with open(log_path, mode='w', newline='', encoding='utf-8') as log_file:
        writer = csv.writer(log_file)
        writer.writerow(["Timestamp", "Filename", "Original", "Replacement"])
        for filename, original, replacement in sorted(unique_redactions):
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            writer.writerow([timestamp, filename, original, replacement])

from docx.oxml.ns import qn

def anonymize_docx_full(input_path, output_path, filename):
    """Anonymize DOCX: paragraphs, tables, headers, footers, and footnotes."""
    doc = Document(input_path)

    # Anonymize paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                run.text = anonymize_text(run.text, filename)

    # Anonymize tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            run.text = anonymize_text(run.text, filename)

    # Anonymize headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    run.text = anonymize_text(run.text, filename)
        for para in section.footer.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    run.text = anonymize_text(run.text, filename)

# Anonymize footnotes - (14/04/25)this is not working yet. Can't identify the footnote elements correctly.
    footnotes = doc.element.xpath(".//w:footnote")
    for footnote in footnotes:
        for text_element in footnote.xpath(".//w:t", namespaces=doc.element.nsmap):
            if text_element.text and text_element.text.strip():
                text_element.text = anonymize_text(text_element.text, filename)
    
    # Anonymize comments
    comments = doc.element.xpath("//w:comment")
    for comment in comments:
        for para in comment.xpath(".//w:p", namespaces=doc.element.nsmap):
            for run in para.xpath(".//w:r", namespaces=doc.element.nsmap):
                for text_element in run.xpath(".//w:t", namespaces=doc.element.nsmap):
                    if text_element.text.strip():
                        text_element.text = anonymize_text(text_element.text, filename)

    # Save the anonymized document
    doc.save(output_path)

def anonymize_excel(input_path, output_path, filename):
    """Anonymize Excel cells using openpyxl while preserving formatting."""
    wb = load_workbook(input_path)
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    original_value = cell.value
                    new_value = anonymize_text(original_value, filename)
                    if new_value != original_value:
                        cell.value = new_value
    wb.save(output_path)

def redact_pdf(input_path, output_path, filename):
    """Redact sensitive entities in a PDF using black boxes."""
    doc = fitz.open(input_path)
    for page in doc:
        text = page.get_text()
        entities = [(ent.text, ent.label_) for ent in nlp(text).ents if ent.label_ in CUSTOM_LABELS]
        for word, label in entities:
            areas = page.search_for(word)
            for area in areas:
                page.add_redact_annot(area, fill=(0, 0, 0))
                unique_redactions.add((filename, word, CUSTOM_LABELS[label]))
        page.apply_redactions()
    doc.save(output_path)

def process_folder(input_folder, output_folder):
    """Process all supported files in the selected folder with progress updates."""
    files = [f for f in os.listdir(input_folder) if os.path.isfile(os.path.join(input_folder, f))]
    total = len(files)

    for i, filename in enumerate(files, start=1):
        name, ext = os.path.splitext(filename.lower())
        input_path = os.path.join(input_folder, filename)
        output_path = os.path.join(output_folder, f"{name}_anonymized{ext}")

        print(f"📄 Processing file {i} of {total}: {filename}")

        try:
            if ext == ".docx":
                anonymize_docx_full(input_path, output_path, filename)
            elif ext == ".xlsx":
                anonymize_excel(input_path, output_path, filename)
            elif ext == ".pdf":
                redact_pdf(input_path, output_path, filename)
            else:
                print(f"⚠️ Skipped unsupported file: {filename}")
        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")

    log_path = os.path.join(output_folder, "anonymization_log.csv")
    save_replacement_log(log_path)
    print(f"📝 Anonymization log saved to: {log_path}")

def process_single_file(input_path, output_folder):
    """Process a single file and save output to the selected folder."""
    filename = os.path.basename(input_path)
    name, ext = os.path.splitext(filename.lower())
    output_path = os.path.join(output_folder, f"{name}_anonymized{ext}")

    print(f"📄 Processing file: {filename}")

    try:
        if ext == ".docx":
            anonymize_docx_full(input_path, output_path, filename)
        elif ext == ".xlsx":
            anonymize_excel(input_path, output_path, filename)
        elif ext == ".pdf":
            redact_pdf(input_path, output_path, filename)
        else:
            print(f"⚠️ Unsupported file type: {filename}")
            return
    except Exception as e:
        print(f"❌ Error processing {filename}: {e}")

    log_path = os.path.join(output_folder, "anonymization_log.csv")
    save_replacement_log(log_path)
    print(f"📝 Anonymization log saved to: {log_path}")

if __name__ == "__main__":
    print("📁 Welcome to the Document Anonymizer Tool")
    print("Choose processing mode:\n1. Single file\n2. Batch folder")
    mode = input("Enter 1 or 2: ").strip()

    # Get the whitelist from the user
    WHITELIST = get_whitelist_from_user()
    print(f"✅ Whitelisted keywords: {', '.join(WHITELIST) if WHITELIST else 'None'}")

    if mode == "1":
        input_file = input("Enter the full path to the file to anonymize: ").strip().strip("'\"")
        output_dir = input("Enter the full path to the output folder: ").strip().strip("'\"")
        if os.path.exists(input_file) and os.path.exists(output_dir):
            unique_redactions.clear()
            entity_mapping.clear()
            process_single_file(input_file, output_dir)
        else:
            print("❌ Invalid file or folder path.")
    elif mode == "2":
        input_dir = input("Enter the full path to the folder with files: ").strip().strip("'\"")
        output_dir = input("Enter the full path to the output folder: ").strip().strip("'\"")
        if os.path.exists(input_dir) and os.path.exists(output_dir):
            unique_redactions.clear()
            entity_mapping.clear()
            process_folder(input_dir, output_dir)
        else:
            print("❌ Invalid folder paths.")
    else:
        print("❗ Invalid selection. Please enter 1 or 2.")
        
# This script is designed to anonymize sensitive information in Word, Excel, and PDF documents.
# End of script